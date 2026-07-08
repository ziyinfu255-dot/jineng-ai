#!/usr/bin/env python3

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, east_asia="宋体", size=12, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="8AA07D", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:color"), color)


def format_body_paragraph(paragraph, indent=True, spacing=1.42):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spacing
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    if indent:
        pf.first_line_indent = Cm(0.74)


def add_title(doc, title):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(title)
    set_run_font(run, east_asia="微软雅黑", size=18, bold=True, color="234B3A")


def add_board_table(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F3E7")
    set_cell_border(cell, color="A8925D", size="10")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, east_asia="等线", size=10.5, color="4A463A")


def build_doc(md_path, out_path, mode):
    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    title = ""
    current_heading = ""
    in_code = False
    board_buffer = []

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            continue

        if not title:
            continue

        if not doc.paragraphs:
            add_title(doc, title)

        if stripped.startswith("```"):
            in_code = not in_code
            if not in_code and board_buffer:
                add_board_table(doc, board_buffer)
                board_buffer = []
            continue

        if in_code:
            if "板书设计" in current_heading:
                board_buffer.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("## "):
            current_heading = stripped[3:].strip()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(current_heading)
            set_run_font(run, east_asia="微软雅黑", size=13.5, bold=True, color="2E5A4A")
            continue

        if stripped.startswith("### "):
            current_heading = stripped[4:].strip()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(current_heading)
            set_run_font(run, east_asia="微软雅黑", size=12, bold=True, color="3D5C7A")
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            pf = paragraph.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.3
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            run = paragraph.add_run(stripped[2:].strip())
            set_run_font(run, east_asia="宋体", size=11.5)
            continue

        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(stripped[2:-2])
            set_run_font(run, east_asia="微软雅黑", size=11.5, bold=True, color="444444")
            continue

        short_centered = (
            mode == "script"
            and len(stripped) <= 22
            and any(token in stripped for token in ("；", "。", "：", "，"))
            and not stripped.startswith(("同学们", "老师", "为什么", "所以", "现在", "下面", "最后", "如果", "请", "来"))
        )
        no_indent = stripped.endswith("：") or stripped.startswith(
            (
                "人人有地块",
                "讲台前排净",
                "卫生角整理",
                "值日分工好",
                "扫帚拖把靠墙站",
                "簸箕挂好不添乱",
                "垃圾入桶套好袋",
                "抹布洗净挂起晒",
                "我承诺",
                "值日生，真光荣",
                "你扫地",
                "教室美",
                "一个班级",
            )
        )

        paragraph = doc.add_paragraph()
        format_body_paragraph(
            paragraph,
            indent=not no_indent,
            spacing=1.42 if mode == "script" else 1.35,
        )
        if short_centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(stripped)
        set_run_font(run, east_asia="宋体", size=12 if mode == "script" else 11.5)

    doc.save(out_path)


def parse_args():
    parser = argparse.ArgumentParser(
        description="Export class-meeting script and lesson-plan Markdown files to styled Word documents."
    )
    parser.add_argument("--script-md", required=True, help="Path to the verbatim script Markdown file")
    parser.add_argument("--lesson-md", required=True, help="Path to the lesson plan Markdown file")
    parser.add_argument("--script-docx", required=True, help="Output path for the verbatim script Word file")
    parser.add_argument("--lesson-docx", required=True, help="Output path for the lesson plan Word file")
    return parser.parse_args()


def main():
    args = parse_args()
    build_doc(args.script_md, args.script_docx, mode="script")
    build_doc(args.lesson_md, args.lesson_docx, mode="lesson")


if __name__ == "__main__":
    main()
